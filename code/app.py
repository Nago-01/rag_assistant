import os, re, traceback
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_groq import ChatGroq
from langchain_google_genai import ChatGoogleGenerativeAI
from datetime import datetime
from .db import VectorDB


# Load environment variables
load_dotenv()

def extract_pdf_metadata(file_path: Path) -> Dict[str, Any]:
    """
    Extract rich metadata from PDF files.

    Returns:
        Dictionary with title, author, date, page_count, etc.
    """
    reader = PdfReader(file_path)
    metadata = {}

    # Extract PyPDF2 metadata
    pdf_info = reader.metadata
    if pdf_info:
        metadata["title"] = pdf_info.get("/Title", file_path.stem)
        metadata["author"] = pdf_info.get("/Author", "Unknown")
        metadata["creation_date"] = pdf_info.get("/CreationDate", "Unknown")
        metadata["subject"] = pdf_info.get("/Subject", " ")
    else:
        metadata["title"] = file_path.stem
        metadata["author"] = "Unknown"

    # Page count
    metadata["page_count"] = len(reader.pages)

    # Extracting date from filename if possible
    date_match = re.search(r'(\d{4})-?(\d{2})?-?(\d{2})?', file_path.name)
    if date_match:
        metadata["date_from_filename"] = date_match.group(0)

    return metadata


def extract_docx_metadata(file_path: Path) -> Dict[str, Any]:
    """Extract metadata from DOCX files."""
    doc = Document(file_path)
    core_props = doc.core_properties

    metadata = {
        "title": core_props.title or file_path.stem,
        "author": core_props.author or "Unknown",
        "creation_date": str(core_props.created) if core_props.created else "Unknown",
        "modified_date": str(core_props.modified) if core_props.modified else "Unknown",
        "subject": core_props.subject or "",
        "page_count": len(doc.sections) 
    }

    return metadata


def extract_text_with_page_numbers(file_path: Path, ext: str) -> List[Dict[str, Any]]:
    """
    Extract text from documents while preserving page numbers.

    Returns:
        List of dictionaries with 'content', 'page_number' and 'section'
    """
    chunks = []

    if ext == ".pdf":
        reader = PdfReader(file_path)
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                chunks.append({
                    "content": text.strip(),
                    "page_number": page_num,
                    "section": f"Page {page_num}"
                })

    elif ext == ".docx":
        doc = Document(file_path)
        current_section = "Introduction"
        page_num = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detecting section headers
            if para.style.name.startswith("Heading"):
                current_section = text

            chunks.append({
                "content": text,
                "page_number": page_num,
                "section": current_section

            })

            # Rough page estimation
            if len(text) > 500:
                page_num += 1

    elif ext in (".txt", ".md"):
        text = file_path.read_text(encoding="utf-8")
        chunks.append({
            "content": text.strip(),
            "page_number": 1,
            "section": "Full Documentation"
        })

    return chunks


# Load documents
def load_publication(pub_dir: Optional[Path] = None) -> List[Dict[str, Any]]:
    """Loads documents with enhanced metadata extraction.

    Returns:
        List of documents as dictionaries with rich metadata.
    """
    # Determine path with multiple checks
    if pub_dir:
        path = Path(pub_dir)
    else:
        # Trying multiple possible paths
        possible_paths = [
            Path("./data"),
            Path(__file__).resolve().parent / "data",
            Path(__file__).resolve().parent.parent / "data",
        ]

        path = None
        for p in possible_paths:
            if p.exists():
                path = p
                print(f"Found data folder at: {path}")
                break

        if path is None:
            error_msg = "Data folder not found. Tried:\n" + "\n".join(f"  - {p.absolute()}" for p in possible_paths)
            raise FileNotFoundError(error_msg)

    if not path.exists():
        raise FileNotFoundError(f"Data folder not found: {path}")
    
    results = []

    try:
        if path.is_file():
            files = [path]
        else:
            # Handle folder containing multiple files
            all_files = list(path.iterdir())
            files = sorted([f for f in all_files if f.suffix.lower() in (".md", ".txt", ".pdf", ".docx")])

        if not files:
            raise FileNotFoundError(f"No supported files found in: {path}")
        
        print(f"Found {len(files)} files: {[file.name for file in files]}\n")
        
        # Process each file
        for idx, file in enumerate(files, 1):
            print(f"\nProcessing file {idx}/{len(files)}: {file.name}")
            ext = file.suffix.lower()

           # Extract document-level metadata
            if ext == ".pdf":
                doc_metadata = extract_pdf_metadata(file)
            elif ext == ".docx":
                doc_metadata = extract_docx_metadata(file)
            else:
                doc_metadata = {
                    "title": file.stem,
                    "author": "Unknown",
                    "page_count": 1
                }

            # Add common metadata
            doc_metadata.update({
                "source": file.name,
                "type": ext.replace(".", ""),
                "path": str(file.resolve()),
                "file_size": file.stat().st_size,
                "last_modified": datetime.fromtimestamp(file.stat().st_mtime).isoformat()
            })

            # Extract text with page numbers
            text_chunks = extract_text_with_page_numbers(file, ext)

            # Creating a result entry for each page
            for chunk in text_chunks:
                results.append({
                    "content": chunk["content"],
                    "metadata": {
                        **doc_metadata, # Document-level metaddata
                        "page_number": chunk["page_number"],
                        "section": chunk["section"]
                    }
                })
        
        print("\n" + "="*50)
        print(f"Successfully loaded {len(results)} document chunks.")
        print("="*50 + "\n")

        return results

    except Exception as e:
        print(f"\nError loading documents: {e}")
        traceback.print_exc()
        raise IOError(f"Error loading documents: {e}") from e


class QAAssistant:
    """
    A simple Hybrid-memory RAG-based QA assistant using ChromaDB and multiple LLM providers.
    Supports Groq and Google Gemini APIs
    """

    def __init__(self):
        """Initialize the RAG assistant with vector DB and LLM."""

        # Initialize LLM and check for valid API in order of preference
        self.llm = self._initialize_llm()
        if not self.llm:
            raise ValueError(
                "No valid API key found."
                "Please set your: GROQ_API_KEY, or GOOGLE_API_KEY in your .env file"
            )
        
        # Initialize vector database
        self.db = VectorDB()

        # Create RAG prompt
        PROMPT = """
        You are a smart, helpful and knowledgeable assistant. Use the provided publication to answer all user's questions.
        If the answer to the question is not contained within the context, respond with "I'm not sure" rather than guessing.
        
        Context:
        {context}

        Question:
        {question}

        Answer in a factual, clear and concise manner.       
        """

        self.prompt_template = ChatPromptTemplate.from_template(PROMPT)

        # Chain the prompt and LLM
        self.chain = self.prompt_template | self.llm | StrOutputParser()

        print("QA Assistant Initialized successfully")

    
    def _initialize_llm(self):
        """
        Initialize the LLM based on available API keys.
        Try Groq first, then Google Gemini.
        """
        # Check for Groq API key
        if os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY"):
            model_name = os.getenv("GEMINI_MODEL", "gemini-2.0-flash-exp")
            return ChatGoogleGenerativeAI(
                google_api_key=os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY"),
                model=model_name,
                temperature=0.0
            )
        elif os.getenv("GROQ_API_KEY"):
            model_name = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
            return ChatGroq(
                api_key=os.getenv("GROQ_API_KEY"),
                model=model_name,
                temperature=0.0
            )
        else:
            return None
    

    def add_doc(self, documents: List) -> None:
        """
        Add documents to the knowledge base.

        Args:
            documents: List of documents
        """
        self.db.add_doc(documents)


    def invoke(self, input: str, n_results: int = 3) -> Dict[str, Any]:
        """
        Query with enhanced citations.

        Returns:
            Dictionary containing answer and citations
        """
        # Retrieve the top relevant chunks from the vector DB
        results = self.db.search(input, n_results, use_reranking=True)

        # Combine the retrieved chunks into a single context string
        context = "\n\n".join(results.get("documents", []))
        citations = results.get("citations", [])

        # Generate the answer using the LLM chain
        llm_answer = self.chain.invoke({"context": context, "question": input})

        # return the answer with citations
        return {
            "answer": llm_answer,
            "citations": citations,
            "sources": list(set([m.get("source") for m in results.get("metadatas", [])]))
        }
    

def main():
    """Main function to demonstrate the QA assistant."""
    pass

if __name__ == "__main__":
    main()