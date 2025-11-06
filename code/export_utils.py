"""
Export utilities for generating LaTex, Word, and PDF documents.
"""

from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime


def generate_latex_bibliography(bibliography: List[Dict[str, Any]], output_path: str = "bibliography.tex") -> str:
    """
    Generate a LaTex bibliography file.

    Args:
        bibliography: List of document metadata dictionaries.
        output_path: Path to save the LaTex file

    Returns:
        Path to the generated file
    """
    latex_content = r"""\documentclass[12pt]{article}
\usepackage[utf8]{inputenc}
\usepackage[margin=1in]{geometry}
\usepackage{hyperref}
\usepackage{natbib}

\title{Bibliography}
\author{Aloysia}
\date{\today}

\begin{document}

\maketitle

\section*{References}

\begin{thebibliography}{99}
"""

    for i, entry in enumerate(bibliography, 1):
        author = entry.get("author", "Unknown Author")
        title = entry.get("title", "Untitled")
        source = entry.get("source", "Unknown Source")
        date = entry.get("date", "n.d.")
        pages = entry.get("page_count", "N/A")

        # Format as LaTex bibliography entry
        latex_content += f"""\\bibitem{{{source.replace('.', '_')}}}
{author}. ({date}). \\textit{{{title}}}. {source}. {pages} pages.

"""
        
        latex_content += r"""\end{thebibliography}
        
\end{document}
"""

        # Save to file
        output = Path(output_path)
        output.write_text(latex_content, encoding='utf-8')

        return str(output.absolute())


def generate_word_bibliography(bibliography: List[Dict[str, Any]], output_path: str = "bibliography.docx") -> str:
    """
    Generate a Word document bibliography.
    

    Args:
        bibliography: List of document metadata dictionaries.
        output_path: Path to save the Word file 

    Returns:
        Path to the generated file
    """
    doc = Document()

    # Title
    title = doc.add_heading('Bibliography', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Total Documents: {len(bibliography)}")
    doc.add_paragraph("")

    # Add heading
    doc.add_heading('References', 1)

    # Add each bibliography entry
    for i, entry in enumerate(bibliography, 1):
        author = entry.get("author", "Unknown Author")
        title = entry.get("title", "Untitled")
        source = entry.get("source", "Unknown Source")
        date = entry.get("date", "n.d.")
        pages = entry.get("page_count", "N/A")
        doc_type = entry.get("type", "Unknown")

        # Create citation paragraph
        citation = doc.add_paragraph()
        citation.add_run(f"{i}. ").bold = True
        citation.add_run(f"{author}. ({date}). ").bold = True
        citation.add_run(f"{title}. ").italic = True
        citation.add_run(f"{source}. {doc_type.upper()}. {pages} pages.")

        # Add spacing
        doc.add_paragraph("")

    # Save document
    output = Path(output_path)
    doc.save(output)

    return str(output.absolute())



def generate_markdown_bibliography(bibliography: List[Dict[str, Any]], output_path: str = "bibliography.md") -> str:
    """
    Generated a Markdown bibliography file.

    Args:
        bibliography: List of document metadata dictionaries.
        output_path: Path to save the Markdown file

    Returns:
        Path to the generated file
    """
    md_content = f"""# Bibliography

**Generated:** {datetime.now().strftime('%B %d, %Y')}
**Total Documents:** {len(bibliography)}

---

##Reference


"""
    
    for i, entry in enumerate(bibliography, 1):
        author = entry.get("author", "Unknown Author")
        title = entry.get("title", "Untitled")
        source = entry.get("source", "Unknown Source")
        date = entry.get("date", "n.d.")
        pages = entry.get("page_count", "N/A")
        doc_type = entry.get("type", "Unknown")

        md_content += f"""
### {i}. {title}


- **Author:** {author}
- **Source:** {source}
- **Type:** {doc_type.upper()}
- **Date:** {date}
- **Pages:** {pages}

"""
        
    # Save to file
    output = Path(output_path)
    output.write_text(md_content, encoding='utf-8')

    return str(output.absolute())



def generate_literature_review_document(
    topic: str,
    sections: List[Dict[str, Any]],
    format: str = "word",
    output_path: str = None
) -> str:
    """
    Generate a complete literature review document.

    Args: 
        topic: The research topic
        sections: List of review sections with content and citations
        format: Output format ('word', 'latex', 'markdown')
        output_path: Path to save the document

    Returns:
        Path to the generated file
    """

    if format == "word":
        return _generate_word_review(topic, sections, output_path or f"literature_review_{topic.replace(' ', '_')}.docx")
    elif format == "latex":
        return _generate_latex_review(topic, sections, output_path or f"literature_review_{topic.replace(' ', '_')}.tex")
    else:
        return _generate_markdown_review(topic, sections, output_path or f"literature_review_{topic.replace(' ', '_')}.md")
    


def _generate_word_review(topic: str, sections: List[Dict[str, Any]], output_path: str) -> str:
    """Generate a Word document for a literature review."""
    doc = Document()

    # Title page
    title = doc.add_heading(f'Literature Review: {topic}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # Abstract
    doc.add_heading('Abstract', 1)
    doc.add_paragraph(
        f"This literature review synthesizes findings from {len(sections)} sources "
        f"related to {topic}. The review examines key themes, methodologies, and "
        f"findings across the included studies."
    )
    doc.add_page_break()

    # Main content sections
    for i, section in enumerate(sections, 1):
        doc.add_heading(f"{i}. {section.get('source', 'Unknown Source')}", 2)

        # Add content
        content = section.get('content', 'No content available')
        doc.add_paragraph(content)

        # Add citations
        citations = section.get('citations', [])
        if citations:
            citation_para = doc.add_paragraph()
            citation_para.add_run("Citations: ").bold = True
            citation_para.add_run("; ".join(citations))

        doc.add_paragraph("")

    # References section
    doc.add_page_break()
    doc.add_heading('References', 1)

    all_citations = []
    for section in sections:
        all_citations.extend(section.get("citations", []))

    for citation in set(all_citations):
        doc.add_paragraph(citation, style="List Number")

    # Save
    output = Path(output_path)
    doc.save(output)
    return str(output.absolute())


def _generate_latex_review(topic: str, sections: List[Dict[str, Any]], output_path: str) -> str:
    """Generate LaTex format literature review."""
    latex_content = r"""\documentclass[12pt]{article}
\usepackage[utf8]{inputenc}
\usepackage[margin=1in]{geometry}
\usepackage{hyperref}
\usepackage{cite}

\title{Literature Review: """ + topic + r"""}
\author{Aloysia}
\date{\today}

\begin{document}

\maketitle

\begin{abstract}
This literature review synthesizes finds from """ + str(len(sections)) + r""" sources related to """ + topic + r""".
The review examines key themes, methodologies, and findings across the included studies.
\end{abstract}

\newpage


"""

    for i, section in enumerate(sections, 1):
        source = section.get("source", "Unknown Source")
        content = section.get("content", "No content available")
        
        latex_content += f"""\\section{{{source}}}

{content}

"""
    latex_content += r"""
\newpage
\section*{References}

"""

    all_citations = []
    for section in sections:
        all_citations.extend(section.get("citations", []))

    for citation in set(all_citations):
        latex_content += f"\\item {citation}\n"

    latex_content += r"""
\end{document}
"""

    output = Path(output_path)
    output.write_text(latex_content, encoding="utf-8")
    return str(output.absolute())


def _generate_markdown_review(topic: str, sections: List[Dict[str, Any]], output_path: str) -> str:
    """Generate Markdown format literature review."""
    md_content = f"""# Literature Review: {topic}
    
**Generated:** {datetime.now().strftime('%B %d, %Y')}

---

##Abstract

This literature review synthesizes findings from {len(sections)} sources related to {topic}.
The review examines key themes, methodologies, and findings across the included studies.

---
    
"""
    
    for i, section in enumerate(sections, 1):
        source = section.get("source", "Unknown Source")
        content = section.get("content", "No content available")
        citations = section.get("citations", [])
        
        md_content += f"""## {i}. {source}

{content}

"""
        if citations:
            md_content += f"**Citations:** {'; '.join(citations)}\n\n"
        
        md_content += "---\n\n"

    # References
    md_content += "##References\n\n"

    all_citations = []
    for section in sections:
        all_citations.extend(section.get("citations", []))

    for i, citation in enumerate(set(all_citations), 1):
        md_content += f"{i}. {citation}\n"


    output = Path(output_path)
    output.write_text(md_content, encoding="utf-8")
    return str(output.absolute())



